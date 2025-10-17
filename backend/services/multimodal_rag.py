# Multi-Modal RAG - Images, Tables, Charts
import logging
from typing import List, Dict, Any, Optional, Tuple
from pathlib import Path
import base64
from io import BytesIO

logger = logging.getLogger(__name__)


class MultiModalProcessor:
    """
    Multi-Modal document processor for RAG.

    Handles:
    - Images (with captions)
    - Tables (structured data)
    - Charts (visual data)
    - Mixed content documents

    Benefits:
    - Complete document understanding
    - Visual information retrieval
    - Structured data queries
    """

    def __init__(
        self, embedding_service, llm_manager, vision_llm_available: bool = False
    ):
        self.embedding_service = embedding_service
        self.llm_manager = llm_manager
        self.vision_llm_available = vision_llm_available

    async def process_document(
        self, file_path: str, document_id: str
    ) -> Dict[str, Any]:
        """
        Process multi-modal document.

        Args:
            file_path: Path to document file
            document_id: Document identifier

        Returns:
            Processed multi-modal content
        """
        try:
            file_extension = Path(file_path).suffix.lower()

            # Extract different content types
            text_chunks = await self._extract_text(file_path, file_extension)
            images = await self._extract_images(file_path, file_extension)
            tables = await self._extract_tables(file_path, file_extension)
            charts = await self._extract_charts(file_path, file_extension)

            # Process each content type
            processed_content = {
                "document_id": document_id,
                "text_chunks": text_chunks,
                "images": [],
                "tables": [],
                "charts": [],
            }

            # Process images
            for img in images:
                processed_img = await self._process_image(img, document_id)
                if processed_img:
                    processed_content["images"].append(processed_img)

            # Process tables
            for table in tables:
                processed_table = await self._process_table(table, document_id)
                if processed_table:
                    processed_content["tables"].append(processed_table)

            # Process charts
            for chart in charts:
                processed_chart = await self._process_chart(chart, document_id)
                if processed_chart:
                    processed_content["charts"].append(processed_chart)

            logger.info(
                f"Processed multi-modal document: "
                f"{len(text_chunks)} text chunks, "
                f"{len(processed_content['images'])} images, "
                f"{len(processed_content['tables'])} tables, "
                f"{len(processed_content['charts'])} charts"
            )

            return processed_content

        except Exception as e:
            logger.error(f"Multi-modal processing failed: {e}")
            return {"document_id": document_id, "error": str(e)}

    async def _extract_text(
        self, file_path: str, file_extension: str
    ) -> List[Dict[str, Any]]:
        """Extract text content"""
        # Use existing document processor
        try:
            if file_extension == ".pdf":
                return await self._extract_text_from_pdf(file_path)
            elif file_extension in [".docx", ".doc"]:
                return await self._extract_text_from_docx(file_path)
            else:
                return []
        except Exception as e:
            logger.debug(f"Text extraction failed: {e}")
            return []

    async def _extract_images(
        self, file_path: str, file_extension: str
    ) -> List[Dict[str, Any]]:
        """Extract images from document"""
        images = []

        try:
            if file_extension == ".pdf":
                # Extract images from PDF
                import fitz  # PyMuPDF

                doc = fitz.open(file_path)

                for page_num in range(len(doc)):
                    page = doc[page_num]
                    image_list = page.get_images()

                    for img_index, img in enumerate(image_list):
                        xref = img[0]
                        base_image = doc.extract_image(xref)

                        images.append(
                            {
                                "page": page_num + 1,
                                "index": img_index,
                                "image_data": base_image["image"],
                                "extension": base_image["ext"],
                            }
                        )

                doc.close()

            elif file_extension in [".docx"]:
                # Extract images from DOCX
                from docx import Document
                from docx.oxml import parse_xml

                doc = Document(file_path)

                for rel in doc.part.rels.values():
                    if "image" in rel.target_ref:
                        images.append(
                            {
                                "image_data": rel.target_part.blob,
                                "extension": rel.target_ref.split(".")[-1],
                            }
                        )

            logger.info(f"Extracted {len(images)} images from {file_path}")

        except Exception as e:
            logger.debug(f"Image extraction failed: {e}")

        return images

    async def _extract_tables(
        self, file_path: str, file_extension: str
    ) -> List[Dict[str, Any]]:
        """Extract tables from document"""
        tables = []

        try:
            if file_extension == ".pdf":
                # Extract tables from PDF using camelot or tabula
                try:
                    import camelot

                    table_list = camelot.read_pdf(file_path, pages="all")

                    for i, table in enumerate(table_list):
                        tables.append(
                            {
                                "index": i,
                                "data": table.df.to_dict("records"),
                                "page": table.page,
                            }
                        )
                except ImportError:
                    logger.debug("Camelot not available for table extraction")

            elif file_extension in [".docx"]:
                # Extract tables from DOCX
                from docx import Document

                doc = Document(file_path)

                for i, table in enumerate(doc.tables):
                    table_data = []
                    for row in table.rows:
                        row_data = [cell.text for cell in row.cells]
                        table_data.append(row_data)

                    tables.append({"index": i, "data": table_data})

            logger.info(f"Extracted {len(tables)} tables from {file_path}")

        except Exception as e:
            logger.debug(f"Table extraction failed: {e}")

        return tables

    async def _extract_charts(
        self, file_path: str, file_extension: str
    ) -> List[Dict[str, Any]]:
        """Extract charts/graphs from document"""
        # Charts are typically embedded as images
        # We'll identify them during image processing
        return []

    async def _process_image(
        self, image_data: Dict[str, Any], document_id: str
    ) -> Optional[Dict[str, Any]]:
        """Process image and generate caption"""
        try:
            # Generate caption using Vision LLM or fallback
            if self.vision_llm_available:
                caption = await self._generate_image_caption_vision(image_data)
            else:
                caption = await self._generate_image_caption_fallback(image_data)

            # Embed caption
            caption_embedding = await self.embedding_service.embed(caption)

            return {
                "type": "image",
                "document_id": document_id,
                "caption": caption,
                "embedding": caption_embedding,
                "metadata": {
                    "page": image_data.get("page"),
                    "index": image_data.get("index"),
                },
            }

        except Exception as e:
            logger.debug(f"Image processing failed: {e}")
            return None

    async def _generate_image_caption_vision(self, image_data: Dict[str, Any]) -> str:
        """Generate caption using Vision LLM"""
        # Placeholder for Vision LLM integration
        # Would use GPT-4 Vision, Claude 3, or similar
        return "Image content (Vision LLM not configured)"

    async def _generate_image_caption_fallback(self, image_data: Dict[str, Any]) -> str:
        """Generate basic caption without Vision LLM"""
        page = image_data.get("page", "unknown")
        index = image_data.get("index", 0)
        return f"Image on page {page}, index {index}"

    async def _process_table(
        self, table_data: Dict[str, Any], document_id: str
    ) -> Optional[Dict[str, Any]]:
        """Process table and generate summary"""
        try:
            # Convert table to text representation
            table_text = self._table_to_text(table_data["data"])

            # Generate summary
            summary = await self._summarize_table(table_text)

            # Embed summary
            summary_embedding = await self.embedding_service.embed(summary)

            return {
                "type": "table",
                "document_id": document_id,
                "summary": summary,
                "full_text": table_text,
                "embedding": summary_embedding,
                "metadata": {
                    "index": table_data.get("index"),
                    "page": table_data.get("page"),
                },
            }

        except Exception as e:
            logger.debug(f"Table processing failed: {e}")
            return None

    def _table_to_text(self, table_data: Any) -> str:
        """Convert table data to text"""
        if isinstance(table_data, list):
            # List of rows
            lines = []
            for row in table_data:
                if isinstance(row, list):
                    lines.append(" | ".join(str(cell) for cell in row))
                elif isinstance(row, dict):
                    lines.append(" | ".join(f"{k}: {v}" for k, v in row.items()))
            return "\n".join(lines)
        else:
            return str(table_data)

    async def _summarize_table(self, table_text: str) -> str:
        """Generate table summary"""
        try:
            prompt = f"""Summarize the following table in 2-3 sentences:

{table_text[:1000]}

Summary:"""

            summary = await self.llm_manager.generate(
                prompt=prompt, max_tokens=150, temperature=0.3
            )

            return summary if summary else table_text[:200]

        except Exception as e:
            logger.debug(f"Table summarization failed: {e}")
            return table_text[:200]

    async def _process_chart(
        self, chart_data: Dict[str, Any], document_id: str
    ) -> Optional[Dict[str, Any]]:
        """Process chart/graph"""
        # Similar to image processing
        return await self._process_image(chart_data, document_id)

    async def _extract_text_from_pdf(self, file_path: str) -> List[Dict[str, Any]]:
        """Extract text from PDF"""
        try:
            import fitz

            doc = fitz.open(file_path)
            chunks = []

            for page_num in range(len(doc)):
                page = doc[page_num]
                text = page.get_text()

                if text.strip():
                    chunks.append({"text": text, "page": page_num + 1})

            doc.close()
            return chunks

        except Exception as e:
            logger.debug(f"PDF text extraction failed: {e}")
            return []

    async def _extract_text_from_docx(self, file_path: str) -> List[Dict[str, Any]]:
        """Extract text from DOCX"""
        try:
            from docx import Document

            doc = Document(file_path)
            chunks = []

            for i, para in enumerate(doc.paragraphs):
                if para.text.strip():
                    chunks.append({"text": para.text, "paragraph": i})

            return chunks

        except Exception as e:
            logger.debug(f"DOCX text extraction failed: {e}")
            return []


class MultiModalRetriever:
    """
    Retriever for multi-modal content.

    Searches across:
    - Text chunks
    - Image captions
    - Table summaries
    - Chart descriptions
    """

    def __init__(self, embedding_service, milvus_manager):
        self.embedding_service = embedding_service
        self.milvus_manager = milvus_manager

    async def search(
        self, query: str, top_k: int = 10, content_types: Optional[List[str]] = None
    ) -> List[Dict[str, Any]]:
        """
        Search across all content types.

        Args:
            query: Search query
            top_k: Number of results
            content_types: Filter by content type (text/image/table/chart)

        Returns:
            Mixed content results
        """
        try:
            # Embed query
            query_embedding = await self.embedding_service.embed(query)

            # Search in Milvus
            results = await self.milvus_manager.search(
                query_vectors=[query_embedding],
                top_k=top_k * 2,
                output_fields=["type", "text", "caption", "summary", "metadata"],
            )

            # Filter by content type if specified
            filtered_results = []
            for result in results[0]:
                content_type = result.get("type", "text")

                if content_types and content_type not in content_types:
                    continue

                filtered_results.append(result)

            return filtered_results[:top_k]

        except Exception as e:
            logger.error(f"Multi-modal search failed: {e}")
            return []


def create_multimodal_processor(
    embedding_service, llm_manager, vision_llm_available: bool = False
) -> MultiModalProcessor:
    """Factory function"""
    return MultiModalProcessor(embedding_service, llm_manager, vision_llm_available)


def create_multimodal_retriever(
    embedding_service, milvus_manager
) -> MultiModalRetriever:
    """Factory function"""
    return MultiModalRetriever(embedding_service, milvus_manager)
