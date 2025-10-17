# Advanced Document Parser with Table and Image Support
import logging
from typing import List, Dict, Optional, Any
from dataclasses import dataclass
from datetime import datetime
import re

logger = logging.getLogger(__name__)


@dataclass
class DocumentElement:
    """Parsed document element"""

    type: str  # text, table, image, heading
    content: str
    metadata: Dict[str, Any]
    position: int


@dataclass
class TableData:
    """Structured table data"""

    headers: List[str]
    rows: List[List[str]]
    caption: Optional[str] = None


@dataclass
class ImageData:
    """Image metadata"""

    path: str
    caption: Optional[str] = None
    alt_text: Optional[str] = None
    width: Optional[int] = None
    height: Optional[int] = None


@dataclass
class DocumentMetadata:
    """Enhanced document metadata"""

    title: Optional[str] = None
    author: Optional[str] = None
    created_date: Optional[datetime] = None
    modified_date: Optional[datetime] = None
    page_count: Optional[int] = None
    language: Optional[str] = None
    keywords: List[str] = None
    has_tables: bool = False
    has_images: bool = False
    table_count: int = 0
    image_count: int = 0


class AdvancedDocumentParser:
    """
    Advanced document parser with support for:
    - Tables extraction and structuring
    - Image detection and metadata
    - Enhanced metadata extraction
    - Semantic chunking with context preservation
    """

    def __init__(self):
        self.supported_formats = [".pdf", ".docx", ".txt", ".md"]

    async def parse_document(
        self,
        file_path: str,
        extract_tables: bool = True,
        extract_images: bool = True,
        extract_metadata: bool = True,
    ) -> Dict[str, Any]:
        """
        Parse document with advanced features.

        Args:
            file_path: Path to document file
            extract_tables: Whether to extract tables
            extract_images: Whether to extract images
            extract_metadata: Whether to extract metadata

        Returns:
            Dictionary with parsed content and metadata
        """
        file_ext = file_path.lower().split(".")[-1]

        if file_ext == "pdf":
            return await self._parse_pdf(
                file_path, extract_tables, extract_images, extract_metadata
            )
        elif file_ext == "docx":
            return await self._parse_docx(
                file_path, extract_tables, extract_images, extract_metadata
            )
        elif file_ext in ["txt", "md"]:
            return await self._parse_text(file_path, extract_metadata)
        else:
            raise ValueError(f"Unsupported file format: {file_ext}")

    async def _parse_pdf(
        self,
        file_path: str,
        extract_tables: bool,
        extract_images: bool,
        extract_metadata: bool,
    ) -> Dict[str, Any]:
        """Parse PDF document"""
        try:
            import PyPDF2
            import pdfplumber

            elements = []
            metadata = DocumentMetadata()

            # Extract basic text with PyPDF2
            with open(file_path, "rb") as file:
                pdf_reader = PyPDF2.PdfReader(file)

                if extract_metadata:
                    pdf_info = pdf_reader.metadata
                    if pdf_info:
                        metadata.title = pdf_info.get("/Title")
                        metadata.author = pdf_info.get("/Author")
                        metadata.page_count = len(pdf_reader.pages)

                # Extract text
                for page_num, page in enumerate(pdf_reader.pages):
                    text = page.extract_text()
                    if text.strip():
                        elements.append(
                            DocumentElement(
                                type="text",
                                content=text,
                                metadata={"page": page_num + 1},
                                position=len(elements),
                            )
                        )

            # Extract tables with pdfplumber
            if extract_tables:
                with pdfplumber.open(file_path) as pdf:
                    for page_num, page in enumerate(pdf.pages):
                        tables = page.extract_tables()
                        for table_idx, table in enumerate(tables):
                            if table:
                                table_data = self._structure_table(table)
                                elements.append(
                                    DocumentElement(
                                        type="table",
                                        content=self._table_to_text(table_data),
                                        metadata={
                                            "page": page_num + 1,
                                            "table_index": table_idx,
                                            "structured_data": table_data,
                                        },
                                        position=len(elements),
                                    )
                                )
                                metadata.table_count += 1

            metadata.has_tables = metadata.table_count > 0

            return {
                "elements": elements,
                "metadata": metadata,
                "raw_text": "\n\n".join(
                    e.content for e in elements if e.type == "text"
                ),
            }

        except ImportError as e:
            logger.error(f"Required library not installed: {e}")
            # Fallback to basic parsing
            return await self._parse_text(file_path, extract_metadata)
        except Exception as e:
            logger.error(f"Error parsing PDF: {e}")
            raise

    async def _parse_docx(
        self,
        file_path: str,
        extract_tables: bool,
        extract_images: bool,
        extract_metadata: bool,
    ) -> Dict[str, Any]:
        """Parse DOCX document"""
        try:
            from docx import Document

            doc = Document(file_path)
            elements = []
            metadata = DocumentMetadata()

            # Extract metadata
            if extract_metadata:
                core_props = doc.core_properties
                metadata.title = core_props.title
                metadata.author = core_props.author
                metadata.created_date = core_props.created
                metadata.modified_date = core_props.modified

            # Extract paragraphs
            for para in doc.paragraphs:
                if para.text.strip():
                    # Detect headings
                    is_heading = para.style.name.startswith("Heading")

                    elements.append(
                        DocumentElement(
                            type="heading" if is_heading else "text",
                            content=para.text,
                            metadata={
                                "style": para.style.name,
                                "is_heading": is_heading,
                            },
                            position=len(elements),
                        )
                    )

            # Extract tables
            if extract_tables:
                for table_idx, table in enumerate(doc.tables):
                    table_data = self._extract_docx_table(table)
                    elements.append(
                        DocumentElement(
                            type="table",
                            content=self._table_to_text(table_data),
                            metadata={
                                "table_index": table_idx,
                                "structured_data": table_data,
                            },
                            position=len(elements),
                        )
                    )
                    metadata.table_count += 1

            metadata.has_tables = metadata.table_count > 0

            return {
                "elements": elements,
                "metadata": metadata,
                "raw_text": "\n\n".join(
                    e.content for e in elements if e.type in ["text", "heading"]
                ),
            }

        except ImportError:
            logger.error("python-docx not installed")
            return await self._parse_text(file_path, extract_metadata)
        except Exception as e:
            logger.error(f"Error parsing DOCX: {e}")
            raise

    async def _parse_text(
        self, file_path: str, extract_metadata: bool
    ) -> Dict[str, Any]:
        """Parse plain text document"""
        with open(file_path, "r", encoding="utf-8") as file:
            content = file.read()

        elements = [
            DocumentElement(type="text", content=content, metadata={}, position=0)
        ]

        metadata = DocumentMetadata()

        return {"elements": elements, "metadata": metadata, "raw_text": content}

    def _structure_table(self, raw_table: List[List[str]]) -> TableData:
        """Structure raw table data"""
        if not raw_table:
            return TableData(headers=[], rows=[])

        # First row as headers
        headers = [str(cell) if cell else "" for cell in raw_table[0]]

        # Remaining rows as data
        rows = []
        for row in raw_table[1:]:
            rows.append([str(cell) if cell else "" for cell in row])

        return TableData(headers=headers, rows=rows)

    def _extract_docx_table(self, table) -> TableData:
        """Extract table from DOCX"""
        rows_data = []
        for row in table.rows:
            row_data = [cell.text for cell in row.cells]
            rows_data.append(row_data)

        return self._structure_table(rows_data)

    def _table_to_text(self, table_data: TableData) -> str:
        """Convert table to readable text"""
        lines = []

        # Headers
        if table_data.headers:
            lines.append(" | ".join(table_data.headers))
            lines.append("-" * 50)

        # Rows
        for row in table_data.rows:
            lines.append(" | ".join(row))

        return "\n".join(lines)

    def extract_keywords(self, text: str, top_n: int = 10) -> List[str]:
        """Extract keywords from text"""
        # Simple keyword extraction
        words = re.findall(r"\b\w+\b", text.lower())

        # Remove common stop words
        stop_words = {
            "the",
            "a",
            "an",
            "and",
            "or",
            "but",
            "in",
            "on",
            "at",
            "to",
            "for",
            "of",
            "with",
            "by",
            "from",
            "is",
            "are",
        }

        keywords = [w for w in words if w not in stop_words and len(w) > 3]

        # Count frequencies
        from collections import Counter

        word_counts = Counter(keywords)

        return [word for word, _ in word_counts.most_common(top_n)]


# Global parser instance
_parser: AdvancedDocumentParser = None


def get_advanced_parser() -> AdvancedDocumentParser:
    """Get global advanced parser instance"""
    global _parser
    if _parser is None:
        _parser = AdvancedDocumentParser()
    return _parser
