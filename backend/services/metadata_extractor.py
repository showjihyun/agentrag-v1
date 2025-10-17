"""
Metadata Extraction Service for documents.

Extracts rich metadata from various document formats including:
- Author, creation date, modification date
- Document title and headings
- Language detection
- Keywords extraction
- Document structure (sections, tables, images)
"""

import logging
import io
from typing import Dict, Any, List, Optional
from datetime import datetime
import re

# Document processing libraries
from PyPDF2 import PdfReader
from docx import Document as DocxDocument
import langdetect

logger = logging.getLogger(__name__)


class MetadataExtractor:
    """
    Service for extracting rich metadata from documents.

    Features:
    - PDF metadata extraction (author, dates, title, etc.)
    - DOCX metadata extraction (core properties)
    - Language detection
    - Keywords extraction
    - Document structure analysis
    """

    def __init__(self):
        """Initialize MetadataExtractor."""
        logger.info("MetadataExtractor initialized")

    def extract_pdf_metadata(
        self, pdf_content: bytes, text: Optional[str] = None
    ) -> Dict[str, Any]:
        """
        Extract metadata from PDF file.

        Args:
            pdf_content: PDF file content as bytes
            text: Extracted text (optional, for language detection)

        Returns:
            Dictionary with PDF metadata
        """
        try:
            pdf_file = io.BytesIO(pdf_content)
            reader = PdfReader(pdf_file)

            # Get PDF metadata
            metadata = reader.metadata if reader.metadata else {}

            # Extract standard fields
            extracted = {
                "title": self._clean_metadata_value(metadata.get("/Title", "")),
                "author": self._clean_metadata_value(metadata.get("/Author", "")),
                "subject": self._clean_metadata_value(metadata.get("/Subject", "")),
                "creator": self._clean_metadata_value(metadata.get("/Creator", "")),
                "producer": self._clean_metadata_value(metadata.get("/Producer", "")),
                "creation_date": self._parse_pdf_date(
                    metadata.get("/CreationDate", "")
                ),
                "modification_date": self._parse_pdf_date(metadata.get("/ModDate", "")),
                "page_count": len(reader.pages),
                "format": "PDF",
            }

            # Add language detection if text provided
            if text:
                extracted["language"] = self._detect_language(text)
                extracted["keywords"] = self._extract_keywords(text)
                extracted["word_count"] = len(text.split())
                extracted["character_count"] = len(text)

            logger.debug(
                f"Extracted PDF metadata: {extracted.get('title', 'Untitled')}"
            )

            return extracted

        except Exception as e:
            logger.error(f"Failed to extract PDF metadata: {e}")
            return {"format": "PDF", "error": str(e)}

    def extract_docx_metadata(
        self, docx_content: bytes, text: Optional[str] = None
    ) -> Dict[str, Any]:
        """
        Extract metadata from DOCX file.

        Args:
            docx_content: DOCX file content as bytes
            text: Extracted text (optional, for additional analysis)

        Returns:
            Dictionary with DOCX metadata
        """
        try:
            docx_file = io.BytesIO(docx_content)
            doc = DocxDocument(docx_file)

            # Get core properties
            core_props = doc.core_properties

            extracted = {
                "title": core_props.title or "",
                "author": core_props.author or "",
                "subject": core_props.subject or "",
                "keywords": core_props.keywords or "",
                "comments": core_props.comments or "",
                "category": core_props.category or "",
                "creation_date": (
                    core_props.created.isoformat() if core_props.created else None
                ),
                "modification_date": (
                    core_props.modified.isoformat() if core_props.modified else None
                ),
                "last_modified_by": core_props.last_modified_by or "",
                "revision": core_props.revision,
                "format": "DOCX",
            }

            # Document structure
            extracted["paragraph_count"] = len(doc.paragraphs)
            extracted["table_count"] = len(doc.tables)
            extracted["image_count"] = len(doc.inline_shapes)

            # Extract headings
            headings = self._extract_docx_headings(doc)
            if headings:
                extracted["headings"] = headings
                extracted["has_toc"] = True

            # Add language detection if text provided
            if text:
                extracted["language"] = self._detect_language(text)
                if not extracted["keywords"]:
                    extracted["keywords"] = ", ".join(self._extract_keywords(text))
                extracted["word_count"] = len(text.split())
                extracted["character_count"] = len(text)

            logger.debug(
                f"Extracted DOCX metadata: {extracted.get('title', 'Untitled')}"
            )

            return extracted

        except Exception as e:
            logger.error(f"Failed to extract DOCX metadata: {e}")
            return {"format": "DOCX", "error": str(e)}

    def extract_txt_metadata(self, txt_content: bytes, text: str) -> Dict[str, Any]:
        """
        Extract metadata from TXT file.

        Args:
            txt_content: TXT file content as bytes
            text: Extracted text

        Returns:
            Dictionary with TXT metadata
        """
        try:
            extracted = {
                "format": "TXT",
                "encoding": self._detect_encoding(txt_content),
                "language": self._detect_language(text),
                "keywords": self._extract_keywords(text),
                "word_count": len(text.split()),
                "character_count": len(text),
                "line_count": text.count("\n") + 1,
            }

            # Try to extract title from first line
            lines = text.split("\n")
            if lines:
                first_line = lines[0].strip()
                if len(first_line) < 100:  # Reasonable title length
                    extracted["title"] = first_line

            logger.debug("Extracted TXT metadata")

            return extracted

        except Exception as e:
            logger.error(f"Failed to extract TXT metadata: {e}")
            return {"format": "TXT", "error": str(e)}

    def extract_hwp_metadata(
        self, hwp_content: bytes, text: Optional[str] = None
    ) -> Dict[str, Any]:
        """
        Extract metadata from HWP file.

        Args:
            hwp_content: HWP file content as bytes
            text: Extracted text (optional)

        Returns:
            Dictionary with HWP metadata
        """
        try:
            import olefile

            hwp_file = io.BytesIO(hwp_content)
            ole = olefile.OleFileIO(hwp_file)

            extracted = {"format": "HWP", "language": "ko"}  # HWP is primarily Korean

            # Try to extract document info
            if ole.exists("DocInfo"):
                # HWP metadata is complex, basic extraction
                extracted["has_doc_info"] = True

            # Count sections
            section_count = 0
            for entry in ole.listdir():
                if "BodyText/Section" in "/".join(entry):
                    section_count += 1
            extracted["section_count"] = section_count

            ole.close()

            # Add text-based metadata if available
            if text:
                extracted["keywords"] = self._extract_keywords(text)
                extracted["word_count"] = len(text.split())
                extracted["character_count"] = len(text)

            logger.debug("Extracted HWP metadata")

            return extracted

        except Exception as e:
            logger.error(f"Failed to extract HWP metadata: {e}")
            return {"format": "HWP", "language": "ko", "error": str(e)}

    def extract_hwpx_metadata(
        self, hwpx_content: bytes, text: Optional[str] = None
    ) -> Dict[str, Any]:
        """
        Extract metadata from HWPX file.

        Args:
            hwpx_content: HWPX file content as bytes
            text: Extracted text (optional)

        Returns:
            Dictionary with HWPX metadata
        """
        try:
            import zipfile
            import xml.etree.ElementTree as ET

            hwpx_file = io.BytesIO(hwpx_content)

            extracted = {"format": "HWPX", "language": "ko"}  # HWPX is primarily Korean

            with zipfile.ZipFile(hwpx_file, "r") as zip_ref:
                # Try to read metadata from docInfo.xml
                if "docInfo.xml" in zip_ref.namelist():
                    doc_info = zip_ref.read("docInfo.xml")
                    root = ET.fromstring(doc_info)

                    # Extract basic info (structure varies by version)
                    extracted["has_doc_info"] = True

                # Count sections
                section_files = [
                    f
                    for f in zip_ref.namelist()
                    if f.startswith("Contents/section") and f.endswith(".xml")
                ]
                extracted["section_count"] = len(section_files)

            # Add text-based metadata if available
            if text:
                extracted["keywords"] = self._extract_keywords(text)
                extracted["word_count"] = len(text.split())
                extracted["character_count"] = len(text)

            logger.debug("Extracted HWPX metadata")

            return extracted

        except Exception as e:
            logger.error(f"Failed to extract HWPX metadata: {e}")
            return {"format": "HWPX", "language": "ko", "error": str(e)}

    def extract_metadata(
        self, file_content: bytes, file_type: str, text: Optional[str] = None
    ) -> Dict[str, Any]:
        """
        Extract metadata based on file type.

        Args:
            file_content: File content as bytes
            file_type: Type of file (pdf, txt, docx, hwp, hwpx)
            text: Extracted text (optional)

        Returns:
            Dictionary with extracted metadata
        """
        extractors = {
            "pdf": self.extract_pdf_metadata,
            "docx": self.extract_docx_metadata,
            "txt": lambda content, txt: self.extract_txt_metadata(content, txt or ""),
            "hwp": self.extract_hwp_metadata,
            "hwpx": self.extract_hwpx_metadata,
        }

        extractor = extractors.get(file_type)
        if not extractor:
            logger.warning(f"No metadata extractor for file type: {file_type}")
            return {"format": file_type.upper()}

        try:
            if file_type == "txt":
                return extractor(file_content, text)
            else:
                return extractor(file_content, text)
        except Exception as e:
            logger.error(f"Metadata extraction failed for {file_type}: {e}")
            return {"format": file_type.upper(), "error": str(e)}

    def _clean_metadata_value(self, value: Any) -> str:
        """Clean and normalize metadata value."""
        if not value:
            return ""

        # Convert to string
        value_str = str(value)

        # Remove null bytes and control characters
        value_str = "".join(char for char in value_str if char.isprintable())

        return value_str.strip()

    def _parse_pdf_date(self, date_str: str) -> Optional[str]:
        """
        Parse PDF date format to ISO format.

        PDF date format: D:YYYYMMDDHHmmSSOHH'mm'
        """
        if not date_str:
            return None

        try:
            # Remove D: prefix if present
            if date_str.startswith("D:"):
                date_str = date_str[2:]

            # Extract date components
            if len(date_str) >= 14:
                year = int(date_str[0:4])
                month = int(date_str[4:6])
                day = int(date_str[6:8])
                hour = int(date_str[8:10])
                minute = int(date_str[10:12])
                second = int(date_str[12:14])

                dt = datetime(year, month, day, hour, minute, second)
                return dt.isoformat()
        except Exception as e:
            logger.debug(f"Failed to parse PDF date '{date_str}': {e}")

        return None

    def _detect_language(self, text: str, sample_size: int = 1000) -> str:
        """
        Detect language of text.

        Args:
            text: Text to analyze
            sample_size: Number of characters to sample

        Returns:
            ISO 639-1 language code (e.g., 'en', 'ko', 'ja')
        """
        if not text or len(text.strip()) < 10:
            return "unknown"

        try:
            # Use sample for efficiency
            sample = text[:sample_size]
            lang = langdetect.detect(sample)
            return lang
        except Exception as e:
            logger.debug(f"Language detection failed: {e}")
            return "unknown"

    def _detect_encoding(self, content: bytes) -> str:
        """Detect text encoding."""
        try:
            import chardet

            result = chardet.detect(content)
            return result["encoding"] or "utf-8"
        except:
            # Fallback: try common encodings
            for encoding in ["utf-8", "utf-8-sig", "latin-1", "cp1252"]:
                try:
                    content.decode(encoding)
                    return encoding
                except:
                    continue
            return "unknown"

    def _extract_keywords(self, text: str, max_keywords: int = 10) -> List[str]:
        """
        Extract keywords from text using simple frequency analysis.

        Args:
            text: Text to analyze
            max_keywords: Maximum number of keywords to return

        Returns:
            List of keywords
        """
        if not text or len(text.strip()) < 50:
            return []

        try:
            # Simple keyword extraction
            # Remove punctuation and convert to lowercase
            words = re.findall(r"\b\w+\b", text.lower())

            # Filter out short words and common stop words
            stop_words = {
                "the",
                "a",
                "an",
                "and",
                "or",
                "but",
                "in",
                "on",
                "at",
                "to",
                "for",
                "of",
                "with",
                "by",
                "from",
                "as",
                "is",
                "was",
                "are",
                "were",
                "be",
                "been",
                "being",
                "have",
                "has",
                "had",
                "do",
                "does",
                "did",
                "will",
                "would",
                "should",
                "could",
                "may",
                "might",
                "must",
                "can",
                "this",
                "that",
                "these",
                "those",
                "i",
                "you",
                "he",
                "she",
                "it",
                "we",
                "they",
                # Korean stop words
                "이",
                "그",
                "저",
                "것",
                "수",
                "등",
                "및",
                "또는",
                "하다",
                "되다",
                "있다",
                "없다",
                "이다",
                "아니다",
                "하는",
                "되는",
                "있는",
                "없는",
            }

            filtered_words = [
                word for word in words if len(word) > 2 and word not in stop_words
            ]

            # Count frequency
            from collections import Counter

            word_freq = Counter(filtered_words)

            # Get top keywords
            keywords = [word for word, _ in word_freq.most_common(max_keywords)]

            return keywords

        except Exception as e:
            logger.debug(f"Keyword extraction failed: {e}")
            return []

    def _extract_docx_headings(self, doc: DocxDocument) -> List[Dict[str, Any]]:
        """
        Extract headings from DOCX document.

        Args:
            doc: DocxDocument object

        Returns:
            List of headings with level and text
        """
        headings = []

        try:
            for para in doc.paragraphs:
                if para.style.name.startswith("Heading"):
                    # Extract heading level
                    level_match = re.search(r"Heading (\d+)", para.style.name)
                    level = int(level_match.group(1)) if level_match else 1

                    headings.append({"level": level, "text": para.text.strip()})

            logger.debug(f"Extracted {len(headings)} headings from DOCX")

        except Exception as e:
            logger.debug(f"Failed to extract headings: {e}")

        return headings


# Singleton instance
_metadata_extractor: Optional[MetadataExtractor] = None


def get_metadata_extractor() -> MetadataExtractor:
    """
    Get or create MetadataExtractor singleton instance.

    Returns:
        MetadataExtractor: Singleton instance
    """
    global _metadata_extractor

    if _metadata_extractor is None:
        _metadata_extractor = MetadataExtractor()

    return _metadata_extractor
